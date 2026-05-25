import sys
import os
import io
import tempfile
import random
import string
import base64
from omniORB import CORBA, PortableServer
import pdfservice, pdfservice__POA

from pdf2docx import Converter
import pytesseract
from PIL import Image
import docx

class ConversionImpl(pdfservice__POA.Conversion):
    def convertPdfToWord(self, pdfFile):
        print("Received PDF to Word conversion request.", flush=True)
        try:
            # Save bytes to a temp PDF file
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as temp_pdf:
                temp_pdf.write(pdfFile)
                temp_pdf_path = temp_pdf.name
            
            # Create a temp DOCX file path
            temp_docx_path = temp_pdf_path.replace(".pdf", ".docx")
            
            # Convert using pdf2docx
            print(f"Converting {temp_pdf_path} to {temp_docx_path}...", flush=True)
            cv = Converter(temp_pdf_path)
            cv.convert(temp_docx_path, start=0, end=None)
            cv.close()
            
            # Read DOCX back
            with open(temp_docx_path, "rb") as f:
                docx_bytes = f.read()
                
            # Cleanup
            os.remove(temp_pdf_path)
            os.remove(temp_docx_path)
            
            print("Conversion successful.", flush=True)
            return docx_bytes
        except Exception as e:
            print(f"Error during PDF to Word conversion: {e}", flush=True)
            return b""

    def performOcr(self, imageFile):
        print("Received OCR request.", flush=True)
        try:
            # Read image from bytes
            image = Image.open(io.BytesIO(imageFile))
            # Perform OCR (French + English if available, or just French)
            print("Running tesseract...", flush=True)
            text = pytesseract.image_to_string(image, lang='fra+eng')
            
            # Create a Word document
            doc = docx.Document()
            doc.add_heading('Résultat OCR', 0)
            doc.add_paragraph(text)
            
            # Save to bytes
            docx_io = io.BytesIO()
            doc.save(docx_io)
            print("OCR successful.", flush=True)
            return docx_io.getvalue()
        except Exception as e:
            print(f"Error during OCR: {e}", flush=True)
            return b""

    def mergePdfs(self, pdfFiles):
        print(f"Received Merge PDF request for {len(pdfFiles)} files.", flush=True)
        try:
            import pypdf
            merger = pypdf.PdfMerger()
            for f_bytes in pdfFiles:
                merger.append(io.BytesIO(f_bytes))
            
            out = io.BytesIO()
            merger.write(out)
            merger.close()
            print("Merge successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Merge: {e}", flush=True)
            return b""

    def splitPdf(self, pdfFile, pagesToKeep):
        print(f"Received Split PDF request to keep pages: {list(pagesToKeep)}", flush=True)
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            for idx in pagesToKeep:
                if 0 <= idx < len(reader.pages):
                    writer.add_page(reader.pages[idx])
            
            out = io.BytesIO()
            writer.write(out)
            print("Split successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Split: {e}", flush=True)
            return b""

    def deletePages(self, pdfFile, pagesToDelete):
        print(f"Received Delete Pages request for pages: {list(pagesToDelete)}", flush=True)
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            to_delete = set(pagesToDelete)
            for i in range(len(reader.pages)):
                # 1-indexed page checking
                if (i + 1) not in to_delete:
                    writer.add_page(reader.pages[i])
                    
            out = io.BytesIO()
            writer.write(out)
            print("Delete pages successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Delete pages: {e}", flush=True)
            return b""

    def rotatePages(self, pdfFile, rotations):
        print(f"Received Rotate Pages request with {len(rotations)} rotation instructions.", flush=True)
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            rot_dict = {entry.pageNum: entry.angle for entry in rotations}
            for i, page in enumerate(reader.pages):
                page_num = i + 1
                if page_num in rot_dict:
                    # In pypdf, rotate page by angle degrees clockwise
                    page.rotate(rot_dict[page_num])
                writer.add_page(page)
                
            out = io.BytesIO()
            writer.write(out)
            print("Rotate pages successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Rotate: {e}", flush=True)
            return b""

    def compressPdf(self, pdfFile):
        print("Received Compress PDF request.", flush=True)
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            for page in reader.pages:
                page.compress_content_streams()
                writer.add_page(page)
                
            out = io.BytesIO()
            writer.write(out)
            print("Compression successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Compression: {e}", flush=True)
            return b""

    def protectPdf(self, pdfFile, options):
        print(f"Received Protect PDF request.", flush=True)
        try:
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            for page in reader.pages:
                writer.add_page(page)
            
            # Use provided password or generate a random one for owner
            user_password = options.password if options.password else "".join(random.choices(string.ascii_letters + string.digits, k=8))
            owner_pass = "".join(random.choices(string.ascii_letters + string.digits, k=12))
            
            # Calculate permission flags (PDF standard bits)
            # Start with all permissions enabled
            perms = -1
            
            # If specific restrictions are requested, calculate flags
            if options.noPrint or options.noEdit or options.noCopy or options.noAnnotate:
                perms = 0
                if not options.noPrint: perms |= 0x4          # Bit 2: Print
                if not options.noEdit: perms |= 0x8           # Bit 3: Modify
                if not options.noCopy: perms |= 0x10          # Bit 4: Copy
                if not options.noAnnotate: perms |= 0x20      # Bit 5: Annotate
                perms |= 0xFFFFFF00                           # Keep all other bits
            
            print(f"Encrypting PDF with user_password={'*' * len(user_password)}, permissions=0x{perms:x}", flush=True)
            writer.encrypt(user_password=user_password, owner_password=owner_pass, permissions_flag=perms)
            
            out = io.BytesIO()
            writer.write(out)
            print("Protect PDF successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Protect: {e}", flush=True)
            import traceback
            traceback.print_exc()
            return b""

    def imageToPdf(self, imageFile, mimeType):
        print(f"Received Image to PDF request (mimeType: {mimeType}).", flush=True)
        try:
            img = Image.open(io.BytesIO(imageFile))
            if img.mode in ("RGBA", "P"):
                img = img.convert("RGB")
                
            out = io.BytesIO()
            img.save(out, format="PDF")
            print("Image to PDF successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Image to PDF: {e}", flush=True)
            return b""

    def signPdf(self, pdfFile, signature):
        print(f"Received Sign PDF request on page {signature.pageNum}.", flush=True)
        try:
            import pypdf
            from reportlab.pdfgen import canvas
            
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            target_page_idx = signature.pageNum - 1
            if target_page_idx < 0 or target_page_idx >= len(reader.pages):
                target_page_idx = 0
                
            for i, page in enumerate(reader.pages):
                if i == target_page_idx:
                    width = float(page.mediabox.width)
                    height = float(page.mediabox.height)
                    
                    sig_w = 150.0
                    sig_h = 60.0
                    
                    # Convert percentage coordinates
                    x_pt = (signature.x / 100.0) * width - (sig_w / 2.0)
                    y_pt = (1.0 - (signature.y / 100.0)) * height - (sig_h / 2.0)
                    
                    # Boundary protection
                    x_pt = max(0.0, min(x_pt, width - sig_w))
                    y_pt = max(0.0, min(y_pt, height - sig_h))
                    
                    packet = io.BytesIO()
                    can = canvas.Canvas(packet, pagesize=(width, height))
                    
                    if signature.type == "type":
                        can.setFont("Helvetica-Oblique", 24)
                        can.setFillColorRGB(0.8, 0.0, 0.0)
                        can.drawString(x_pt + 10, y_pt + 20, signature.content)
                    elif signature.type == "draw":
                        header, encoded = signature.content.split(",", 1)
                        img_data = base64.b64decode(encoded)
                        
                        with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as temp_img:
                            temp_img.write(img_data)
                            temp_img_path = temp_img.name
                            
                        can.drawImage(temp_img_path, x_pt, y_pt, width=sig_w, height=sig_h, mask='auto')
                        os.remove(temp_img_path)
                        
                    can.save()
                    packet.seek(0)
                    
                    sig_reader = pypdf.PdfReader(packet)
                    page.merge_page(sig_reader.pages[0])
                    
                writer.add_page(page)
                
            out = io.BytesIO()
            writer.write(out)
            print("Sign PDF successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Sign: {e}", flush=True)
            return b""

    def watermarkPdf(self, pdfFile, watermark):
        print(f"Received Watermark PDF request. text={watermark.text}", flush=True)
        try:
            import pypdf
            from reportlab.pdfgen import canvas
            from reportlab.lib.colors import black

            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()

            for i, page in enumerate(reader.pages):
                width = float(page.mediabox.width)
                height = float(page.mediabox.height)

                packet = io.BytesIO()
                can = canvas.Canvas(packet, pagesize=(width, height))

                # watermark properties
                text = watermark.text if watermark.text else ""
                opacity = float(watermark.opacity) if hasattr(watermark, 'opacity') else 0.15
                angle = float(watermark.angle) if hasattr(watermark, 'angle') else 45.0
                fontSize = float(watermark.fontSize) if hasattr(watermark, 'fontSize') else 60
                x_pct = float(watermark.x) if hasattr(watermark, 'x') else 50
                y_pct = float(watermark.y) if hasattr(watermark, 'y') else 50
                all_pages = bool(watermark.allPages) if hasattr(watermark, 'allPages') else True

                # If not all pages and this is not the first page, just merge original
                if (not all_pages) and (i != 0):
                    writer.add_page(page)
                    continue

                # Calculate position
                x_pt = (x_pct / 100.0) * width
                y_pt = (1.0 - (y_pct / 100.0)) * height

                can.saveState()
                can.translate(x_pt, y_pt)
                can.rotate(angle)
                can.setFillColorRGB(0.6, 0.6, 0.6)
                try:
                    can.setFont("Helvetica", fontSize)
                except Exception:
                    pass
                # ReportLab doesn't support alpha directly here; emulate with gray color
                can.drawString(-fontSize * (len(text) / 4), 0, text)
                can.restoreState()
                can.save()
                packet.seek(0)

                wm_reader = pypdf.PdfReader(packet)
                try:
                    page.merge_page(wm_reader.pages[0])
                except Exception:
                    # fallback: add original page without merging
                    pass

                writer.add_page(page)

            out = io.BytesIO()
            writer.write(out)
            print("Watermark successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Watermark: {e}", flush=True)
            import traceback
            traceback.print_exc()
            return b""

    def createForm(self, pdfFile, fields, targetPage):
        print(f"Received Create Form request with {len(fields)} fields on page {targetPage}.", flush=True)
        try:
            import pypdf
            from reportlab.pdfgen import canvas
            from reportlab.lib.colors import black, red
            
            reader = pypdf.PdfReader(io.BytesIO(pdfFile))
            writer = pypdf.PdfWriter()
            
            target_page_idx = targetPage - 1
            if target_page_idx < 0 or target_page_idx >= len(reader.pages):
                target_page_idx = 0
                
            for i, page in enumerate(reader.pages):
                if i == target_page_idx:
                    width = float(page.mediabox.width)
                    height = float(page.mediabox.height)
                    
                    packet = io.BytesIO()
                    can = canvas.Canvas(packet, pagesize=(width, height))
                    
                    for idx, field in enumerate(fields):
                        x_pt = (field.x / 100.0) * width
                        y_pt = (1.0 - (field.y / 100.0)) * height
                        
                        # Draw form field boxes
                        if field.fieldType == "text":
                            # Draw a text field rectangle
                            can.rect(x_pt - 50.0, y_pt - 8.0, 100.0, 16.0)
                            can.setFont("Helvetica", 10)
                            can.drawString(x_pt - 40.0, y_pt - 2.0, field.val[:20] if field.val else field.label)
                        elif field.fieldType == "checkbox":
                            # Draw a checkbox
                            can.rect(x_pt - 7.5, y_pt - 7.5, 15.0, 15.0)
                            if field.val == "checked":
                                can.setFont("Helvetica", 12)
                                can.drawString(x_pt - 4.0, y_pt - 3.0, "X")
                            
                    can.save()
                    packet.seek(0)
                    
                    form_reader = pypdf.PdfReader(packet)
                    page.merge_page(form_reader.pages[0])
                    
                writer.add_page(page)
                
            out = io.BytesIO()
            writer.write(out)
            print("Create Form successful.", flush=True)
            return out.getvalue()
        except Exception as e:
            print(f"Error during Create Form: {e}", flush=True)
            import traceback
            traceback.print_exc()
            return b""

def main():
    # Initialize ORB
    orb = CORBA.ORB_init(sys.argv, CORBA.ORB_ID)
    
    # Get RootPOA
    poa = orb.resolve_initial_references("RootPOA")
    poaManager = poa._get_the_POAManager()
    poaManager.activate()
    
    # Create implementation
    servant = ConversionImpl()
    
    # Register to Naming Service
    try:
        obj = orb.resolve_initial_references("NameService")
        import CosNaming
        rootContext = obj._narrow(CosNaming.NamingContext)
        
        name = [CosNaming.NameComponent("PdfManager", "")]
        rootContext.rebind(name, servant._this())
        print("Successfully bound to NameService as PdfManager.", flush=True)
    except Exception as e:
        print(f"Failed to bind to NameService: {e}", flush=True)
        sys.exit(1)
        
    print("CORBA Conversion Server is running...", flush=True)
    orb.run()

if __name__ == "__main__":
    main()
